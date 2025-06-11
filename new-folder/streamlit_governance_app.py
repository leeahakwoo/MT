
import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from io import BytesIO

st.set_page_config(page_title="AI 거버넌스 자동 생성기", layout="wide")
st.title("🤖 AI 거버넌스 문서 자동 생성기 (ISO/IEC 42001 기반)")

# --- Sidebar 고급 기능 ---
st.sidebar.header("🛠️ 고급 기능")
risk_suggestion = st.sidebar.checkbox("🔍 리스크 자동 제안")
sensitivity_check = st.sidebar.checkbox("🔒 민감도 분류 지원")
dashboard_preview = st.sidebar.checkbox("📊 대시보드 미리보기")

# --- 입력 폼 시작 ---
st.subheader("1. 조직의 맥락 및 역할")
context = st.text_area("✅ 조직의 외부/내부 환경 이슈를 기술해주세요:",
                       placeholder="예: 인력 부족, 평가 편차, 기술 변화 등")
role = st.selectbox("✅ 조직의 AI 역할을 선택해주세요:",
                    ["개발자", "제공자", "운영자", "사용자", "복합적 역할"])

st.subheader("2. 이해관계자")
stakeholders = st.multiselect("✅ 주요 이해관계자를 선택하세요:",
                               ["고객", "직원", "규제기관", "사회", "협력사"])
needs = st.text_area("✅ 이해관계자의 요구사항을 간략히 서술해주세요:")

st.subheader("3. 데이터 정보")
data_source = st.text_area("✅ 데이터 출처 및 특성을 기술해주세요:")
data_type = st.radio("✅ 데이터 유형을 선택하세요:",
                     ["정형", "비정형", "민감정보 포함", "공공데이터", "혼합형"])

st.subheader("4. 내부 정책 및 시스템")
policy_input = st.text_area("✅ 존재하는 내부 AI 정책 또는 관련 규정을 입력해주세요:")
infrastructure = st.text_area("✅ AI 시스템 인프라(서버, 로깅, 백업 등)를 간략히 기술해주세요:")

st.subheader("5. 책임 및 역할 분장")
cto_name = st.text_input("CTO (총괄 책임자) 이름")
tech_team_role = st.text_area("기술팀 역할 및 책임:")
quality_team_role = st.text_area("품질팀 역할 및 책임:")

# --- 리스크 자동 제안 ---
if risk_suggestion:
    st.subheader("🔍 제안된 리스크 항목")
    st.markdown("- 데이터 편향 및 대표성 부족")
    st.markdown("- 설명 불가능한 결과 또는 자동화 오류")
    st.markdown("- 사용자 오용 또는 오해 가능성")
    st.markdown("- 법/규정 위반 가능성")

# --- 민감도 분류 지원 ---
if sensitivity_check:
    st.subheader("🔒 예상 민감도 결과")
    if data_type and "민감정보" in data_type:
        st.error("❗ 예상 민감도: 높음")
    else:
        st.success("✅ 예상 민감도: 보통 또는 낮음")

# --- 대시보드 요약 ---
if dashboard_preview:
    st.subheader("📊 입력 정보 요약 대시보드")
    summary_df = pd.DataFrame({
        "항목": ["조직 역할", "이해관계자 수", "데이터 유형", "리스크 제안 여부"],
        "내용": [role, len(stakeholders), data_type, "ON" if risk_suggestion else "OFF"]
    })
    st.table(summary_df)

# --- 규칙 기반 문장 생성 함수 ---
def generate_governance_paragraphs(context, role, stakeholders, needs, data_source, data_type,
                                   policy_input, infrastructure, cto_name, tech_team_role, quality_team_role):
    p = []
    p.append(f"당사는 {context} 등의 환경 문제 해결을 위해 AI 기반 품질컨설팅 에이전트를 운영합니다.")
    p.append(f"조직은 '{role}'의 역할을 중심으로 AI 시스템을 기획 및 실행하고 있습니다.")
    if stakeholders:
        p.append(f"주요 이해관계자는 {', '.join(stakeholders)}이며, 이들은 '{needs}'를 요구합니다.")
    else:
        p.append(f"이해관계자의 요구사항은 다음과 같습니다: {needs}")
    p.append(f"데이터는 '{data_source}'에서 수집된 '{data_type}' 유형의 데이터를 기반으로 분석됩니다.")
    p.append(f"현재 적용 중인 내부 정책은 다음과 같습니다: {policy_input}")
    p.append(f"AI 인프라는 다음 요소로 구성됩니다: {infrastructure}")
    p.append(f"CTO {cto_name}는 정책 및 시스템의 총괄 책임을 지며, 기술팀은 '{tech_team_role}', 품질팀은 '{quality_team_role}' 역할을 수행합니다.")
    return "\n\n".join(p)

# --- 문서 생성 함수 (메모리 기반) ---
def generate_docx_in_memory():
    doc = Document()
    doc.add_heading("AI 거버넌스 문서 (ISO/IEC 42001 기반)", 0)

    summary_paragraph = generate_governance_paragraphs(context, role, stakeholders, needs,
                                                       data_source, data_type, policy_input,
                                                       infrastructure, cto_name, tech_team_role, quality_team_role)
    doc.add_paragraph(summary_paragraph)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 문서 생성 버튼 ---
st.markdown("---")
if st.button("📄 문서 생성하기"):
    docx_buffer = generate_docx_in_memory()
    filename = f"AI_Governance_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    st.download_button(
        label="📥 문서 다운로드 (Word)",
        data=docx_buffer,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
